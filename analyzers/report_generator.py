"""
报告生成器 - 将分析结果生成为格式化的审计报告
"""
import json
import os
from datetime import datetime, timedelta
from typing import Dict, Any, Optional, List
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ReportGenerator:
    """
    报告生成器类
    """
    
    def __init__(self, analysis_result_path: str = None, data: dict = None):
        """
        初始化报告生成器
        :param analysis_result_path: 分析结果JSON文件路径
        :param data: 直接传入的分析结果数据（字典格式）
        """
        if data is not None:
            self.data = data
        else:
            self.analysis_result_path = analysis_result_path
            self.data = self._load_analysis_result()
    
    def _load_analysis_result(self) -> Dict[str, Any]:
        """加载分析结果数据"""
        try:
            with open(self.analysis_result_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"加载分析结果失败: {e}")
            return {}
    
    def _get_week_info(self, analysis_id: str) -> dict:
        """从分析ID提取周信息"""
        # 分析ID格式: YYYYMMDDHHMMSS
        if len(analysis_id) >= 8:
            date_str = analysis_id[:8]
            date = datetime.strptime(date_str, '%Y%m%d')
            
            # 计算年份和月份
            year = date.year
            month = date.month
            
            # 计算是当月第几周（周一为一周开始）
            first_day = date.replace(day=1)
            first_day_weekday = first_day.weekday()  # 0=周一, 6=周日
            
            # 计算第一周开始日期（如果1号不是周一，第一周从1号开始）
            if first_day_weekday == 0:
                first_monday = first_day
            else:
                first_monday = first_day + timedelta(days=(7 - first_day_weekday))
            
            # 计算当前日期所在周
            days_since_first_monday = (date - first_monday).days
            week_number = (days_since_first_monday // 7) + 1
            
            # 如果第一周不完整（少于4天），算作第0周，下周开始算第1周
            if first_day_weekday > 3 and week_number == 1:
                # 需要判断是否应该属于上月最后一周
                prev_month_last_day = first_day - timedelta(days=1)
                prev_month_last_weekday = prev_month_last_day.weekday()
                if prev_month_last_weekday >= 4:  # 周四及以后
                    week_number = 1
                else:
                    week_number = 1  # 按自然周计算
            
            return {
                'year': year,
                'month': month,
                'week': week_number,
                'date': date
            }
        return {'year': 2026, 'month': 5, 'week': 1, 'date': datetime.now()}
    
    def _generate_time_range(self, week_info: dict) -> str:
        """生成时间范围字符串"""
        date = week_info['date']
        
        # 获取周一日期
        monday = date - timedelta(days=date.weekday())
        
        # 获取周日日期
        sunday = monday + timedelta(days=6)
        
        # 格式化输出
        return f"{monday.year}年{monday.month}月{monday.day}日00:00:00至{sunday.year}年{sunday.month}月{sunday.day}日23:59:59"
    
    def _generate_workload_section(self) -> str:
        """生成工作量部分"""
        workload = self.data.get('full_result', {}).get('workload', {})
        result = workload.get('result', [])
        
        if not result:
            return "暂无工作量数据"
        
        # 按小组汇总
        group_summary = {}
        for item in result:
            group = item.get('所属小组', '未分组')
            if group not in group_summary:
                group_summary[group] = {
                    '定档量': 0,
                    '已发布': 0,
                    '媒介人数': 0
                }
            group_summary[group]['定档量'] += item.get('定档量', 0)
            group_summary[group]['已发布'] += item.get('已发布', 0)
            group_summary[group]['媒介人数'] += 1
        
        section = "点击图片可查看完整电子表格\n\n"
        
        # 添加汇总表格
        section += "| 所属小组 | 媒介人数 | 定档量 | 已发布 |\n"
        section += "| --- | --- | --- | --- |\n"
        
        for group, stats in group_summary.items():
            section += f"| {group} | {stats['媒介人数']} | {stats['定档量']} | {stats['已发布']} |\n"
        
        # 总计
        total_media = sum(s['媒介人数'] for s in group_summary.values())
        total_scheduled = sum(s['定档量'] for s in group_summary.values())
        total_published = sum(s['已发布'] for s in group_summary.values())
        section += f"| **合计** | **{total_media}** | **{total_scheduled}** | **{total_published}** |\n"
        
        return section
    
    def _generate_quality_section(self) -> str:
        """生成工作质量部分"""
        quality = self.data.get('full_result', {}).get('quality', {})
        group_summary = quality.get('group_summary', [])
        
        if not group_summary:
            return "暂无工作质量数据"
        
        section = "注：由于各组的表头字段不统一，且与数据库字段不匹配，因此过筛滤以最终通过的过筛为准。\n\n"
        
        # 优质达人
        section += "#### （1）优质达人\n\n"
        section += "点击图片可查看完整电子表格\n\n"
        
        # 添加小组汇总表格
        section += "| 所属小组 | 媒介数量 | 总提报达人数 | 过筛人数 | 小组过筛率(%) | 过筛率中位数(%) | 优秀媒介数 | 良好媒介数 |\n"
        section += "| --- | --- | --- | --- | --- | --- | --- | --- |\n"
        
        for item in group_summary:
            section += f"| {item.get('所属小组', '')} | {item.get('媒介数量', 0)} | {item.get('总提报达人数', 0)} | {item.get('总过筛人数', 0)} | {item.get('小组过筛率(%)', '0%')} | {item.get('过筛率中位数(%)', '0%')} | {item.get('优秀媒介数', 0)} | {item.get('良好媒介数', 0)} |\n"
        
        section += "\n"
        
        # 高阅读达人（如果有数据）
        premium_detail = quality.get('premium_detail', [])
        if premium_detail:
            section += "#### （2）高阅读达人\n\n"
            section += "点击图片可查看完整电子表格\n\n"
        else:
            section += "#### （2）高阅读达人\n\n"
            section += "暂无高阅读达人数据\n\n"
        
        # 质量评估说明
        section += "注：\n"
        section += "• 优秀：80%及以上\n"
        section += "• 良好：65%-80%\n"
        section += "• 一般：50%-64%\n"
        section += "• 待改进：40%-49%\n"
        section += "• 较差：低于40%\n\n"
        
        return section
    
    def _generate_cost_section(self) -> str:
        """生成成本发挥部分"""
        cost = self.data.get('full_result', {}).get('cost', {})
        
        section = "注：不含高返、高阅读、晨光超品日\n\n"
        section += "点击图片可查看完整电子表格\n\n"
        
        # 返点分析
        section += "#### （1）返点分析\n\n"
        
        # 返点整体效益分析
        section += "#### ①返点整体效益分析\n\n"
        section += "点击图片可查看完整电子表格\n\n"
        
        # 基于达人量级分析
        section += "#### ②基于达人量级分析\n\n"
        section += "点击图片可查看完整电子表格\n\n"
        
        # 返点表现评估说明
        section += "注：返点表现评估的范围为：\n"
        section += "• 很差：小于10%\n"
        section += "• 较差：10%-20%\n"
        section += "• 一般：20%-25%\n"
        section += "• 良好：25%-35%\n"
        section += "• 优秀：大于35%\n\n"
        
        # 成本分析
        section += "#### （2）成本分析\n\n"
        section += "点击图片可查看完整电子表格\n\n"
        
        return section
    
    def generate_report(self) -> str:
        """生成完整报告"""
        if not self.data:
            return "无法加载分析结果数据"
        
        analysis_id = self.data.get('analysis_id', '')
        week_info = self._get_week_info(analysis_id)
        
        # 标题
        report = f"# 媒介-审计报告（{week_info['year']}年{week_info['month']}月第{week_info['week']}周）\n\n"
        
        # 时间范围说明
        time_range = self._generate_time_range(week_info)
        report += f"以下数据依托于时间为：\n{time_range}数据\n\n"
        
        # 一、总体分析
        report += "## 一、总体分析\n\n"
        
        # 1、总结
        report += "### 1、总结\n\n"
        report += "注：签框量从下周开始统计\n\n"
        report += "点击图片可查看完整电子表格\n\n"
        
        # 2、工作量
        report += "### 2、工作量\n\n"
        report += self._generate_workload_section()
        report += "\n"
        
        # 3、工作质量
        report += "### 3、工作质量\n\n"
        report += self._generate_quality_section()
        report += "\n"
        
        # 4、成本发挥
        report += "### 4、成本发挥\n\n"
        report += self._generate_cost_section()
        
        return report
    
    def save_report(self, output_dir: str = None) -> str:
        """
        保存报告到文件
        :param output_dir: 输出目录，默认为outputs/reports/
        :return: 保存的文件路径
        """
        if output_dir is None:
            output_dir = os.path.join(os.path.dirname(__file__), '..', 'outputs', 'reports')
        
        # 确保目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成文件名
        analysis_id = self.data.get('analysis_id', datetime.now().strftime('%Y%m%d%H%M%S'))
        filename = f"媒介-审计报告_{analysis_id}.md"
        filepath = os.path.join(output_dir, filename)
        
        # 生成并保存报告
        report = self.generate_report()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(report)
        
        return filepath
    
    def generate_word_report(self) -> BytesIO:
        """
        生成 Word 格式报告
        :return: BytesIO 对象，包含 Word 文档内容
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 库未安装，请先安装: pip install python-docx")
        
        if not self.data:
            raise ValueError("无法加载分析结果数据")
        
        doc = Document()
        
        # 设置默认字体为宋体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(11)
        
        analysis_id = self.data.get('analysis_id', '')
        week_info = self._get_week_info(analysis_id)
        time_range = self._generate_time_range(week_info)
        
        # 标题
        title = doc.add_heading(f"媒介-审计报告（{week_info['year']}年{week_info['month']}月第{week_info['week']}周）", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 时间范围说明
        time_para = doc.add_paragraph(f"以下数据依托于时间为：{time_range}数据")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.space_after = Pt(12)
        
        # 一、总体分析
        doc.add_heading('一、总体分析', level=1)
        
        # 1、总结
        doc.add_heading('1、总结', level=2)
        doc.add_paragraph('注：签框量从下周开始统计')
        doc.add_paragraph('点击图片可查看完整电子表格')
        
        # 2、工作量
        doc.add_heading('2、工作量', level=2)
        doc.add_paragraph('点击图片可查看完整电子表格')
        self._add_workload_table_to_word(doc)
        
        # 3、工作质量
        doc.add_heading('3、工作质量', level=2)
        doc.add_paragraph('注：由于各组的表头字段不统一，且与数据库字段不匹配，因此过筛滤以最终通过的过筛为准。')
        self._add_quality_tables_to_word(doc)
        
        # 4、成本发挥
        doc.add_heading('4、成本发挥', level=2)
        doc.add_paragraph('注：不含高返、高阅读、晨光超品日')
        doc.add_paragraph('点击图片可查看完整电子表格')
        self._add_cost_tables_to_word(doc)
        
        # 保存到 BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def _add_workload_table_to_word(self, doc):
        """向 Word 文档添加工作量表格"""
        workload = self.data.get('full_result', {}).get('workload', {})
        result = workload.get('result', [])
        
        if not result:
            doc.add_paragraph('暂无工作量数据')
            return
        
        # 按小组汇总
        group_summary = {}
        for item in result:
            group = item.get('所属小组', '未分组')
            if group not in group_summary:
                group_summary[group] = {
                    '定档量': 0,
                    '已发布': 0,
                    '媒介人数': 0
                }
            group_summary[group]['定档量'] += item.get('定档量', 0)
            group_summary[group]['已发布'] += item.get('已发布', 0)
            group_summary[group]['媒介人数'] += 1
        
        # 创建表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '所属小组'
        hdr_cells[1].text = '媒介人数'
        hdr_cells[2].text = '定档量'
        hdr_cells[3].text = '已发布'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].bold = True
        
        # 添加数据行
        total_media = 0
        total_scheduled = 0
        total_published = 0
        
        for group, stats in group_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = group
            row_cells[1].text = str(stats['媒介人数'])
            row_cells[2].text = str(stats['定档量'])
            row_cells[3].text = str(stats['已发布'])
            
            # 设置单元格样式
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            total_media += stats['媒介人数']
            total_scheduled += stats['定档量']
            total_published += stats['已发布']
        
        # 添加合计行
        row_cells = table.add_row().cells
        row_cells[0].text = '合计'
        row_cells[1].text = str(total_media)
        row_cells[2].text = str(total_scheduled)
        row_cells[3].text = str(total_published)
        
        # 设置合计行样式
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].bold = True
        
        # 添加空段落
        doc.add_paragraph()
    
    def _add_quality_tables_to_word(self, doc):
        """向 Word 文档添加工作质量表格"""
        quality = self.data.get('full_result', {}).get('quality', {})
        group_summary = quality.get('group_summary', [])
        
        if not group_summary:
            doc.add_paragraph('暂无工作质量数据')
            return
        
        # （1）优质达人
        doc.add_heading('（1）优质达人', level=3)
        doc.add_paragraph('点击图片可查看完整电子表格')
        
        # 创建表格
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        headers = ['所属小组', '媒介数量', '总提报达人数', '过筛人数', '小组过筛率(%)', '过筛率中位数(%)', '优秀媒介数', '良好媒介数']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # 设置表头样式
        for cell in hdr_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].bold = True
        
        # 添加数据行
        for item in group_summary:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('所属小组', '')
            row_cells[1].text = str(item.get('媒介数量', 0))
            row_cells[2].text = str(item.get('总提报达人数', 0))
            row_cells[3].text = str(item.get('总过筛人数', 0))
            row_cells[4].text = item.get('小组过筛率(%)', '0%')
            row_cells[5].text = item.get('过筛率中位数(%)', '0%')
            row_cells[6].text = str(item.get('优秀媒介数', 0))
            row_cells[7].text = str(item.get('良好媒介数', 0))
            
            # 设置单元格样式
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空段落
        doc.add_paragraph()
        
        # （2）高阅读达人
        doc.add_heading('（2）高阅读达人', level=3)
        premium_detail = quality.get('premium_detail', [])
        if premium_detail:
            doc.add_paragraph('点击图片可查看完整电子表格')
        else:
            doc.add_paragraph('暂无高阅读达人数据')
        
        # 质量评估说明
        doc.add_paragraph('注：')
        doc.add_paragraph('• 优秀：80%及以上')
        doc.add_paragraph('• 良好：65%-80%')
        doc.add_paragraph('• 一般：50%-64%')
        doc.add_paragraph('• 待改进：40%-49%')
        doc.add_paragraph('• 较差：低于40%')
        
        # 添加空段落
        doc.add_paragraph()
    
    def _add_cost_tables_to_word(self, doc):
        """向 Word 文档添加成本分析表格"""
        # （1）返点分析
        doc.add_heading('（1）返点分析', level=3)
        
        # ①返点整体效益分析
        doc.add_heading('①返点整体效益分析', level=4)
        doc.add_paragraph('点击图片可查看完整电子表格')
        
        # ②基于达人量级分析
        doc.add_heading('②基于达人量级分析', level=4)
        doc.add_paragraph('点击图片可查看完整电子表格')
        
        # 返点表现评估说明
        doc.add_paragraph('注：返点表现评估的范围为：')
        doc.add_paragraph('• 很差：小于10%')
        doc.add_paragraph('• 较差：10%-20%')
        doc.add_paragraph('• 一般：20%-25%')
        doc.add_paragraph('• 良好：25%-35%')
        doc.add_paragraph('• 优秀：大于35%')
        
        # 添加空段落
        doc.add_paragraph()
        
        # （2）成本分析
        doc.add_heading('（2）成本分析', level=3)
        doc.add_paragraph('点击图片可查看完整电子表格')

def generate_report_from_file(analysis_result_path: str, output_dir: str = None) -> str:
    """
    从分析结果文件生成报告
    :param analysis_result_path: 分析结果JSON文件路径
    :param output_dir: 输出目录
    :return: 保存的文件路径
    """
    generator = ReportGenerator(analysis_result_path)
    return generator.save_report(output_dir)

if __name__ == '__main__':
    # 示例用法
    import sys
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        # 默认使用最新的分析结果
        outputs_dir = os.path.join(os.path.dirname(__file__), '..', 'outputs', 'analysis_results')
        files = [f for f in os.listdir(outputs_dir) if f.endswith('.json')]
        if files:
            input_file = os.path.join(outputs_dir, sorted(files)[-1])
        else:
            print("未找到分析结果文件")
            sys.exit(1)
    
    print(f"正在从 {input_file} 生成报告...")
    output_path = generate_report_from_file(input_file)
    print(f"报告已保存到: {output_path}")
