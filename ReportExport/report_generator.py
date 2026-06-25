import json
import os
import logging
import base64
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    HAS_MPL = True
except ImportError:
    HAS_MPL = False
    plt = None

logger = logging.getLogger(__name__)


def load_analysis_data(analysis_id, source='analysis_results'):
    base_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs', source)
    fpath = os.path.join(base_dir, f'{analysis_id}.json')
    if not os.path.exists(fpath):
        raise FileNotFoundError(f"分析结果文件不存在: {fpath}")
    with open(fpath, 'r', encoding='utf-8') as f:
        return json.load(f)


def calc_week_label(dt=None):
    if dt is None:
        dt = datetime.now()
    yesterday = dt
    if dt.date() == datetime.now().date():
        from datetime import timedelta
        yesterday = dt - timedelta(days=1)
    year = yesterday.year
    month = yesterday.month
    week_num = (yesterday.day - 1) // 7 + 1
    return f"{year}年{month:02d}月第{week_num}周"


def calc_month_label(dt=None):
    if dt is None:
        dt = datetime.now()
    return f"{dt.year}年{dt.month:02d}月"


def _safe_val(lst, idx, default=0):
    """按索引安全取值"""
    if not lst:
        return default
    try:
        v = list(lst.values())[idx] if isinstance(lst, dict) else lst[idx]
        if v is None or v == '':
            return default
        # 尝试转为数字
        try:
            return float(v)
        except (ValueError, TypeError):
            return v
    except (IndexError, KeyError, TypeError):
        return default


def _safe_str(v, default=''):
    if v is None or v == '':
        return default
    return str(v)


def build_summary_table_data(data):
    """从分析数据中提取表格数据"""
    full = data.get('full_result', {})
    selected_groups = data.get('selected_groups', [])

    workload = full.get('workload', {})
    quality = full.get('quality', {})
    cost = full.get('cost', {})

    wl_gs = workload.get('group_summary', [])
    cg_gs = cost.get('group_summary', [])
    ql_gs = quality.get('group_summary', [])

    # 工作量: [所属小组, 媒介人数, 总定档量, 定档量占比(%), 已分配, 未分配]
    # 成本: [所属媒介小组, 媒介人数, 定档数量, 总成本, 总报价, 总返点, 平均成本, 平均返点]
    # 质量: [所属小组, 媒介人数, 总提交数量, 提交占比(%), 总共筛数量, 小组筛除(%), 过筛率中位数(%), 优秀媒介数, 良好媒介数, 待改善占比(%)]

    group_map = {}

    for row in wl_gs:
        if isinstance(row, dict):
            vals = list(row.values())
        else:
            continue
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        group_map.setdefault(name, {})
        group_map[name]['media_count'] = _safe_val(vals, 1, 0)
        group_map[name]['dingdang'] = _safe_val(vals, 2, 0)

    for row in cg_gs:
        if isinstance(row, dict):
            vals = list(row.values())
        else:
            continue
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        group_map.setdefault(name, {})
        group_map[name]['avg_cost'] = _safe_val(vals, 6, 0)

    for row in ql_gs:
        if isinstance(row, dict):
            vals = list(row.values())
        else:
            continue
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        group_map.setdefault(name, {})
        group_map[name]['tibao'] = _safe_val(vals, 2, 0)
        guoshuai_raw = vals[6] if len(vals) > 6 else '-'
        if isinstance(guoshuai_raw, (int, float)):
            group_map[name]['guoshuai'] = f"{float(guoshuai_raw):.2f}%"
        else:
            group_map[name]['guoshuai'] = _safe_str(guoshuai_raw, '-')

    table_rows = []
    all_mc = 0
    all_kou = 0
    all_shu = 0
    all_tibao = 0
    all_dingdang = 0
    group_count = 0

    for g in selected_groups:
        info = group_map.get(g, {})
        mc = int(float(info.get('media_count', 0) or 0))
        kou = int(info.get('kou', 0) or 0)
        shu = int(info.get('shu', 0) or 0)
        tibao = int(float(info.get('tibao', 0) or 0))
        dingdang = int(float(info.get('dingdang', 0) or 0))
        avg_dd = round(dingdang / mc, 0) if mc > 0 else 0
        gs_val = info.get('guoshuai', '-')
        ac_val = float(info.get('avg_cost', 0) or 0)

        table_rows.append({
            'group': g,
            'media_count': mc,
            'kou': kou,
            'shu': shu,
            'tibao': tibao,
            'dingdang': dingdang,
            'avg_dingdang': int(avg_dd),
            'guoshuai': gs_val,
            'avg_cost': round(ac_val, 2)
        })

        all_mc += mc
        all_kou += kou
        all_shu += shu
        all_tibao += tibao
        all_dingdang += dingdang
        group_count += 1

    if table_rows:
        avg_all_dd = round(all_dingdang / len(table_rows), 0) if table_rows else 0
        gs_vals = []
        for r in table_rows:
            try:
                gs_vals.append(float(r['guoshuai'].replace('%', '')))
            except (ValueError, AttributeError):
                pass
        avg_gs = f"{sum(gs_vals)/len(gs_vals):.2f}%" if gs_vals else '-'
        avg_ac = round(sum(r['avg_cost'] for r in table_rows) / len(table_rows), 2) if table_rows else 0

        table_rows.append({
            'group': '所有小组',
            'media_count': all_mc,
            'kou': all_kou,
            'shu': all_shu,
            'tibao': all_tibao,
            'dingdang': all_dingdang,
            'avg_dingdang': int(avg_all_dd),
            'guoshuai': avg_gs,
            'avg_cost': avg_ac
        })

    return table_rows


def build_workload_detail_data(data):
    """提取工作量明细数据（按小组分组，每人提报数+定档数）"""
    full = data.get('full_result', {})
    selected_groups = data.get('selected_groups', [])

    quality_result = full.get('quality', {}).get('result', [])
    workload_result = full.get('workload', {}).get('result', [])

    # quality: [媒介ID, 媒介名称, 所属小组, 总提交数量, 共筛除数量, 筛除率(%), 定档量, 主要状态分布]
    # workload: [媒介名称, 所属小组, 定档量, 已分配, 未分配, 综合评级]

    ql_map = {}  # (group, name) → tibao
    for row in quality_result:
        if isinstance(row, dict):
            vals = list(row.values())
        else:
            continue
        name = _safe_str(vals[1] if len(vals) > 1 else '')
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        tibao = _safe_val(vals, 3, 0)
        if name and group:
            ql_map[(group, name)] = int(float(tibao)) if isinstance(tibao, (int, float)) else 0

    wl_map = {}  # (group, name) → dingdang
    for row in workload_result:
        if isinstance(row, dict):
            vals = list(row.values())
        else:
            continue
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        group = _safe_str(vals[1] if len(vals) > 1 else '')
        dingdang = _safe_val(vals, 2, 0)
        if name and group:
            wl_map[(group, name)] = int(float(dingdang)) if isinstance(dingdang, (int, float)) else 0

    # 合并
    all_keys = set(ql_map) | set(wl_map)
    person_map = {}
    for key in all_keys:
        person_map.setdefault(key[0], []).append({
            'name': key[1],
            'tibao': ql_map.get(key, 0),
            'dingdang': wl_map.get(key, 0)
        })

    # 只保留选中小组，并按小组优先级排序
    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}
    sorted_groups = sorted([g for g in selected_groups if g in person_map],
                           key=lambda g: (priority.get(g, 99), g))

    result = []
    for g in sorted_groups:
        people = sorted(person_map[g], key=lambda p: (p['name'] or ''))
        result.append({
            'group': g,
            'people': people
        })

    return result


def _quality_eval_label(rate_str):
    """根据过筛率返回质量评估标签"""
    try:
        rate = float(str(rate_str).replace('%', ''))
    except (ValueError, TypeError):
        return '-'
    if rate >= 80:
        return '优秀'
    elif rate >= 65:
        return '良好'
    elif rate >= 50:
        return '一般'
    elif rate >= 40:
        return '待改进'
    else:
        return '较差'


def _extract_quality_people(records):
    """从质量明细记录列表中提取标准化的 person 列表"""
    # 字段: [媒介ID, 对应名字, 所属小组, 总提报数量, 过筛数量, 过筛率(小数), 过筛率(%), 质量评估, 主要状态分布, 分析类型]
    people = []
    for row in records:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[1] if len(vals) > 1 else '')
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        total = int(float(_safe_val(vals, 3, 0)))
        passed = int(float(_safe_val(vals, 4, 0)))
        rate_raw = vals[5] if len(vals) > 5 else '-'
        rate_str = _safe_str(vals[6] if len(vals) > 6 else '')
        evaluation = _safe_str(vals[7] if len(vals) > 7 else '-')
        if not name:
            continue

        # 统一 rate_decimal / rate_str
        if isinstance(rate_raw, str) and '%' in rate_raw:
            # rate_raw 是 "100.0%" 格式 → 这是过筛率(%)，小数部分从 rate_str 或计算
            rate_decimal = float(rate_raw.replace('%', ''))
            if not rate_str:
                rate_str = rate_raw
        elif isinstance(rate_raw, (int, float)):
            rate_decimal = float(rate_raw)
            if not rate_str:
                rate_str = '%.2f%%' % rate_decimal
        else:
            rate_decimal = 0.0
            if not rate_str:
                rate_str = '0.00%'
        people.append({
            'name': name,
            'group': group,
            'total': total,
            'passed': passed,
            'rate_str': rate_str,
            'rate_decimal': rate_decimal,
            'evaluation': evaluation
        })
    return people


def build_quality_detail_data(data, kind='premium'):
    """提取质量明细数据
    kind: 'premium' → 优质达人 (premium_detail)
          'high_read' → 高阅读达人 (high_read_detail)
    """
    full = data.get('full_result', {})
    qual = full.get('quality', {})

    # 优先从 premium_detail / high_read_detail 取
    detail_key = 'premium_detail' if kind == 'premium' else 'high_read_detail'
    detail = qual.get(detail_key, [])

    if isinstance(detail, list) and detail:
        return _extract_quality_people(detail)

    # fallback: 从 result 中根据分析类型筛选
    quality_result = qual.get('result', [])
    filtered = []
    for row in quality_result:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        purpose = _safe_str(vals[-1] if len(vals) > 8 else '')
        target = '优质达人' if kind == 'premium' else '高阅读达人'
        if purpose == target:
            filtered.append(row)
    if filtered:
        return _extract_quality_people(filtered)

    # 没有对应数据时返回空，不降级到全量 result
    return []


def build_cost_performance_data(data):
    """提取成本发挥数据（按媒介组聚合）
    数据来源：media_group_workload（定档数、平均返点率）+ fixed_media_rebate（最高/最低返点）
    """
    full = data.get('full_result', {})
    cost = full.get('cost', {})
    selected = data.get('selected_groups', [])

    mwg = cost.get('media_group_workload', [])
    reb = cost.get('fixed_media_rebate', [])

    # media_group_workload: [媒介小组, 总达人数, 总项目数, 媒介人数, 已发布数, 未发布数,
    #                        发布率(%), 总成本(元), 平均成本(元), 总下单价(元), 平均下单价(元),
    #                        总返点金额(元), 平均返点金额(元), 平均返点比例(%), 返点表现]
    wl_map = {}
    for row in mwg:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        wl_map[name] = {
            'dingdang': int(float(_safe_val(vals, 1, 0))),
            'total_project': int(float(_safe_val(vals, 2, 0))),
            'avg_cost': float(_safe_val(vals, 8, 0)),
            'avg_order': float(_safe_val(vals, 10, 0)),
            'avg_rebate': _safe_str(vals[13] if len(vals) > 13 else '')
        }

    # fixed_media_rebate: [固定媒介, 定档达人数, 所属小组, 平均返点比例(%), ...,
    #                       总下单价(元), 平均下单价(元), 返点表现评估, 返点优化建议]
    rebate_map = {}
    for row in reb:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        if not group:
            continue
        max_r = vals[4] if len(vals) > 4 else '0%'
        min_r = vals[5] if len(vals) > 5 else '0%'
        rebate_map.setdefault(group, []).append((max_r, min_r))

    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}
    sorted_groups = sorted(selected, key=lambda g: (priority.get(g, 99), g))

    rows = []
    for g in sorted_groups:
        wl = wl_map.get(g, {})
        dingdang = wl.get('dingdang', 0)
        total_project = wl.get('total_project', 0)
        avg_cost = wl.get('avg_cost', 0)
        avg_order = wl.get('avg_order', 0)
        avg_rebate = wl.get('avg_rebate', '-')

        group_rebates = rebate_map.get(g, [])
        if group_rebates:
            max_vals = []
            min_vals = []
            for mr, nr in group_rebates:
                try:
                    max_vals.append(float(str(mr).replace('%', '')))
                    min_vals.append(float(str(nr).replace('%', '')))
                except (ValueError, TypeError):
                    pass
            high = '%.2f%%' % max(max_vals) if max_vals else '-'
            low = '%.2f%%' % min(min_vals) if min_vals else '-'
        else:
            high = '-'
            low = '-'

        # format avg_rebate consistently
        try:
            avg_fmt = '%.2f%%' % float(str(avg_rebate).replace('%', ''))
        except (ValueError, TypeError):
            avg_fmt = str(avg_rebate)

        rows.append({
            'group': g,
            'dingdang': dingdang,
            'total_project': total_project,
            'avg_cost': avg_cost,
            'avg_order': avg_order,
            'avg_rebate': avg_fmt,
            'high': high,
            'low': low
        })

    return rows


def build_level_analysis_data(data):
    """提取基于达人量级分析数据
    数据源: cost.fixed_media_level
    字段: [固定媒介, 达人量级, 达人数, 所属小组, 总成本(元), 平均成本(元),
           总下单价(元), 平均下单价(元), 总返点金额(元), 平均返点金额(元),
           平均返点比例(%), 总互动量, 平均互动量, 总阅读量, 平均阅读量,
           平均CPE, 平均CPM]
    """
    full = data.get('full_result', {})
    lvl = full.get('cost', {}).get('fixed_media_level', [])
    selected = data.get('selected_groups', [])

    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}

    rows = []
    for row in lvl:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        level = _safe_str(vals[1] if len(vals) > 1 else '')
        cnt = int(float(_safe_val(vals, 2, 0)))
        group = _safe_str(vals[3] if len(vals) > 3 else '')
        if not name or group not in selected:
            continue
        rows.append({
            'name': name,
            'level': level,
            'count': cnt,
            'group': group,
            'total_cost': float(_safe_val(vals, 4, 0)),
            'avg_cost': float(_safe_val(vals, 5, 0)),
            'total_order': float(_safe_val(vals, 6, 0)),
            'avg_order': float(_safe_val(vals, 7, 0)),
            'total_rebate': float(_safe_val(vals, 8, 0)),
            'avg_rebate': float(_safe_val(vals, 9, 0)),
            'avg_rebate_pct': _safe_str(vals[10] if len(vals) > 10 else '-'),
            'total_interact': float(_safe_val(vals, 11, 0)),
            'avg_interact': float(_safe_val(vals, 12, 0)),
            'total_read': float(_safe_val(vals, 13, 0)),
            'avg_read': float(_safe_val(vals, 14, 0)),
            'avg_cpe': float(_safe_val(vals, 15, 0)),
            'avg_cpm': float(_safe_val(vals, 16, 0))
        })

    rows.sort(key=lambda r: (priority.get(r['group'], 99), r['group'], r['name'] or '', r['level'] or ''))
    return rows


def build_cost_analysis_data(data):
    """提取定档媒介成本分析数据
    数据源: cost.fixed_media_cost
    字段: [固定媒介, 定档数量, 所属小组, 总成本(元), 平均成本(元), 成本最大值(元),
           成本最小值(元), 成本中位数(元), 总报价(元), 平均报价(元), 报价最大值(元),
           报价最小值(元), 总下单价(元), 平均下单价(元), 总节约金额(元),
           平均节约金额(元), 成本占比(%), 总返点金额(元), 平均返点金额(元)]
    """
    full = data.get('full_result', {})
    fc = full.get('cost', {}).get('fixed_media_cost', [])
    selected = data.get('selected_groups', [])

    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}
    rows = []
    for row in fc:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        if not name or group not in selected:
            continue
        rows.append({
            'name': name,
            'count': int(float(_safe_val(vals, 1, 0))),
            'group': group,
            'total_cost': float(_safe_val(vals, 3, 0)),
            'avg_cost': float(_safe_val(vals, 4, 0)),
            'max_cost': float(_safe_val(vals, 5, 0)),
            'min_cost': float(_safe_val(vals, 6, 0)),
            'median_cost': float(_safe_val(vals, 7, 0)),
            'total_quote': float(_safe_val(vals, 8, 0)),
            'avg_quote': float(_safe_val(vals, 9, 0)),
            'max_quote': float(_safe_val(vals, 10, 0)),
            'min_quote': float(_safe_val(vals, 11, 0)),
            'total_order': float(_safe_val(vals, 12, 0)),
            'avg_order': float(_safe_val(vals, 13, 0)),
            'total_save': float(_safe_val(vals, 14, 0)),
            'avg_save': float(_safe_val(vals, 15, 0)),
            'cost_pct': _safe_str(vals[16] if len(vals) > 16 else '-'),
            'total_rebate': float(_safe_val(vals, 17, 0)),
            'avg_rebate': float(_safe_val(vals, 18, 0))
        })
    rows.sort(key=lambda r: (priority.get(r['group'], 99), r['group'], r['name'] or ''))
    return rows


GROUP_COLORS = {
    '耐消媒介组': ('#F9A825', '#FFF8E1', '#FFECB3'),
    '快消媒介组': ('#2d6a4f', '#d4edda', '#e8f5e9'),
    'AI媒介组':  ('#4361ee', '#dce6ff', '#f0f4ff'),
}
GROUP_KEYS_WITH_PRIORITY = ['耐消媒介组', '快消媒介组', 'AI媒介组']


def _set_cell_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    tcPr.append(shading)


def _set_alt_row_shading(table, row_idx, start_col=0, col_count=None, light='#f8f9fa'):
    """为行设置浅色交替背景"""
    if row_idx % 2 == 0:
        return
    total = len(table.rows[row_idx].cells)
    cols = col_count or (total - start_col)
    end = min(start_col + cols, total)
    for ci in range(start_col, end):
        c = table.rows[row_idx].cells[ci]
        if c:
            _set_cell_shading(c, light)


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(*[int(h[i:i+2], 16) for i in (0, 2, 4)])


def _add_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=None, font_name='微软雅黑'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    if tcPr.find(qn('w:noWrap')) is None:
        tcPr.append(OxmlElement('w:noWrap'))


def _try_get_font():
    """尝试找到支持中文的字体"""
    if not HAS_MPL:
        return None
    candidates = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS', 'DejaVu Sans']
    for name in candidates:
        try:
            prop = fm.FontProperties(family=name)
            if prop.get_name():
                return name
        except Exception:
            continue
    # fallback — find any ttf that has CJK
    try:
        for f in fm.findSystemFonts():
            if any(k in f.lower() for k in ['yahei', 'simhei', 'noto', 'cjk']):
                return fm.FontProperties(fname=f).get_name()
    except Exception:
        pass
    return 'sans-serif'


def build_chart_data(data):
    """提取图表数据"""
    full = data.get('full_result', {})
    selected = data.get('selected_groups', [])

    wl_gs = full.get('workload', {}).get('group_summary', [])
    ql_gs = full.get('quality', {}).get('group_summary', [])

    # quality: [小组, 人数, 总提交量, 占比, 总筛除量, 筛除率%, 过筛率中位数%, ...]
    # workload: [小组, 人数, 总定档量, 占比, 已分配, 未分配]

    ql_map = {}
    for row in ql_gs:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        tibao = _safe_val(vals, 2, 0)
        shai = _safe_val(vals, 4, 0)
        guoshuai = _safe_val(vals, 6, '-')
        ql_map[name] = {'tibao': int(float(tibao)), 'shai': int(float(shai)), 'guoshuai': guoshuai}

    wl_map = {}
    for row in wl_gs:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        if not name:
            continue
        dingdang = _safe_val(vals, 2, 0)
        media_cnt = _safe_val(vals, 1, 0)
        wl_map[name] = {'dingdang': int(float(dingdang)), 'media_count': int(float(media_cnt))}

    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}
    sorted_groups = sorted(selected, key=lambda g: (priority.get(g, 99), g))

    chart_rows = []
    for g in sorted_groups:
        q = ql_map.get(g, {})
        w = wl_map.get(g, {})
        guoshuai_raw = q.get('guoshuai', '-')
        if isinstance(guoshuai_raw, (int, float)):
            guoshuai_pct = float(guoshuai_raw)
        else:
            try:
                guoshuai_pct = float(str(guoshuai_raw).replace('%', ''))
            except (ValueError, TypeError):
                guoshuai_pct = None
        chart_rows.append({
            'group': g,
            'tibao': q.get('tibao', 0),
            'dingdang': w.get('dingdang', 0),
            'media_count': w.get('media_count', 0),
            'guoshuai': guoshuai_pct
        })

    return chart_rows


def _barh_with_labels(ax, y_pos, values, colors, xlabel, title, fmt='%.0f', font_name=None):
    bars = ax.barh(y_pos, values, color=colors, height=0.6, edgecolor='white', linewidth=0.5)
    for bar, v in zip(bars, values):
        label = fmt % v
        ax.text(bar.get_width() + max(values) * 0.015, bar.get_y() + bar.get_height() / 2,
                label, va='center', fontsize=9, fontfamily=font_name)
    ax.set_xlabel(xlabel, fontsize=11, fontfamily=font_name)
    ax.set_title(title, fontsize=13, fontweight='bold', fontfamily=font_name)
    ax.tick_params(axis='y', labelsize=10)
    ax.tick_params(axis='x', labelsize=9)
    for label in ax.get_yticklabels():
        label.set_fontproperties(fm.FontProperties(family=font_name)) if font_name else None
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)


def generate_chart_images(chart_data, output_dir):
    """生成图表并保存为 PNG，返回 {过筛率: path, 定档量: path}"""
    paths = {}
    if not HAS_MPL or not chart_data:
        return paths

    font_name = _try_get_font()
    if font_name:
        plt.rcParams['font.family'] = font_name
    plt.rcParams['axes.unicode_minus'] = False

    groups = [r['group'] for r in chart_data]

    # ── 图1: 小组过筛率 ──
    fig, ax = plt.subplots(figsize=(8, 4))
    guoshuai_vals = []
    guoshuai_colors = []
    for r in chart_data:
        v = r['guoshuai']
        if v is None:
            v = 0
        guoshuai_vals.append(v)
        guoshuai_colors.append('#4ecdc4' if v and v >= 10 else '#ff6b6b')

    y_pos = list(range(len(groups)))
    _barh_with_labels(ax, y_pos, guoshuai_vals, guoshuai_colors,
                      '过筛率 (%)', '各小组过筛率中位数', fmt='%.1f%%', font_name=font_name)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(groups)
    # 标注均值线
    valid = [v for v in guoshuai_vals if v is not None]
    if valid:
        avg = sum(valid) / len(valid)
        ax.axvline(avg, color='#4361ee', linestyle='--', linewidth=1, alpha=0.7)
        ax.text(avg + max(guoshuai_vals) * 0.02, -0.4, '均值 %.1f%%' % avg,
                fontsize=9, color='#4361ee', fontfamily=font_name)
    fig.tight_layout()
    p1 = os.path.join(output_dir, '_chart_guoshuai.png')
    fig.savefig(p1, dpi=150, bbox_inches='tight')
    plt.close(fig)
    paths['guoshuai'] = p1

    # ── 图2: 小组定档量 ──
    fig2, ax2 = plt.subplots(figsize=(8, 4))
    dingdang_vals = [r['dingdang'] for r in chart_data]
    dd_colors = ['#4361ee'] * len(dingdang_vals)
    _barh_with_labels(ax2, list(range(len(groups))), dingdang_vals, dd_colors,
                      '定档量 (条)', '各小组总定档量', fmt='%.0f', font_name=font_name)
    ax2.set_yticks(list(range(len(groups))))
    ax2.set_yticklabels(groups)
    # 标注每个组的媒介人数
    for i, r in enumerate(chart_data):
        mc = r['media_count']
        v = r['dingdang']
        if mc > 0:
            ax2.text(v + max(dingdang_vals) * 0.02, i - 0.2, '人均: %d' % round(v / mc),
                     fontsize=8, color='#666', fontfamily=font_name)
    fig2.tight_layout()
    p2 = os.path.join(output_dir, '_chart_dingdang.png')
    fig2.savefig(p2, dpi=150, bbox_inches='tight')
    plt.close(fig2)
    paths['dingdang'] = p2

    return paths


def build_rebate_detail_data(data):
    """提取每人返点比例数据（用于图表）"""
    full = data.get('full_result', {})
    reb = full.get('cost', {}).get('fixed_media_rebate', [])
    people = []
    for row in reb:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        rate_raw = vals[3] if len(vals) > 3 else '0%'
        dingdang = int(float(_safe_val(vals, 1, 0)))
        if not name:
            continue
        try:
            rate = float(str(rate_raw).replace('%', ''))
        except (ValueError, TypeError):
            rate = 0
        people.append({
            'name': name,
            'group': group,
            'rate': rate,
            'dingdang': dingdang
        })
    return people


def build_rebate_analysis_data(data):
    """提取返点整体效益分析表数据
    数据源: cost.fixed_media_rebate
    字段: [固定媒介, 定档数量, 所属小组, 平均返点率(%), 返点率最大值(%), 返点率最小值(%),
           返点率中位数(%), 返点分布, 总返点金额(元), 平均返点金额(元), 返点金额最大值(元),
           返点金额最小值(元), 返点金额中位数(元), 总下单价(元), 平均下单价(元), 评估, 建议]
    """
    full = data.get('full_result', {})
    reb = full.get('cost', {}).get('fixed_media_rebate', [])
    selected = data.get('selected_groups', [])

    priority = {'耐消媒介组': 0, '快消媒介组': 1, 'AI媒介组': 2}

    rows = []
    for row in reb:
        if not isinstance(row, dict):
            continue
        vals = list(row.values())
        name = _safe_str(vals[0] if len(vals) > 0 else '')
        group = _safe_str(vals[2] if len(vals) > 2 else '')
        if not name or group not in selected:
            continue
        rows.append({
            'name': name,
            'group': group,
            'dingdang': int(float(_safe_val(vals, 1, 0))),
            'avg_rebate': _safe_str(vals[3] if len(vals) > 3 else '-'),
            'max_rebate': _safe_str(vals[4] if len(vals) > 4 else '-'),
            'min_rebate': _safe_str(vals[5] if len(vals) > 5 else '-'),
            'median_rebate': _safe_str(vals[6] if len(vals) > 6 else '-'),
            'total_rebate_amt': float(_safe_val(vals, 8, 0)),
            'avg_rebate_amt': float(_safe_val(vals, 9, 0)),
            'max_rebate_amt': float(_safe_val(vals, 10, 0)),
            'min_rebate_amt': float(_safe_val(vals, 11, 0)),
            'median_rebate_amt': float(_safe_val(vals, 12, 0)),
            'total_order': float(_safe_val(vals, 13, 0)),
            'avg_order': float(_safe_val(vals, 14, 0)),
            'evaluation': _safe_str(vals[15] if len(vals) > 15 else '-'),
            'suggestion': _safe_str(vals[16] if len(vals) > 16 else '-')
        })

    # 按小组优先级排序，组内按名字
    rows.sort(key=lambda r: (priority.get(r['group'], 99), r['group'], r['name'] or ''))
    return rows


def generate_rebate_chart_images(data, output_dir):
    """生成返点比例分布图（高 >35%，低 <20%）"""
    paths = {}
    if not HAS_MPL:
        return paths

    font_name = _try_get_font()
    if font_name:
        plt.rcParams['font.family'] = font_name
    plt.rcParams['axes.unicode_minus'] = False

    people = build_rebate_detail_data(data)
    if not people:
        return paths

    high = [p for p in people if p['rate'] > 35]
    low = [p for p in people if p['rate'] < 20]

    colors_high = ['#ff6b6b' if r['rate'] > 50 else '#ffa94d' for r in high]
    colors_low = ['#4ecdc4' if r['rate'] >= 10 else '#868e96' for r in low]

    # 高返点
    if high:
        high_sorted = sorted(high, key=lambda p: p['rate'])
        names_h = ['%s(%s)' % (p['name'], p['group']) for p in high_sorted]
        vals_h = [p['rate'] for p in high_sorted]
        fig, ax = plt.subplots(figsize=(8, max(3, len(high) * 0.35)))
        _barh_with_labels(ax, list(range(len(vals_h))), vals_h, colors_high,
                          '平均返点比例 (%)', '高返点媒介（>35%）', fmt='%.1f%%', font_name=font_name)
        ax.set_yticks(list(range(len(vals_h))))
        ax.set_yticklabels(names_h)
        fig.tight_layout()
        p_h = os.path.join(output_dir, '_chart_rebate_high.png')
        fig.savefig(p_h, dpi=150, bbox_inches='tight')
        plt.close(fig)
        paths['rebate_high'] = p_h

    # 低返点
    if low:
        low_sorted = sorted(low, key=lambda p: p['rate'])
        names_l = ['%s(%s)' % (p['name'], p['group']) for p in low_sorted]
        vals_l = [p['rate'] for p in low_sorted]
        fig2, ax2 = plt.subplots(figsize=(8, max(3, len(low) * 0.35)))
        _barh_with_labels(ax2, list(range(len(vals_l))), vals_l, colors_low,
                          '平均返点比例 (%)', '低返点媒介（<20%）', fmt='%.1f%%', font_name=font_name)
        ax2.set_yticks(list(range(len(vals_l))))
        ax2.set_yticklabels(names_l)
        fig2.tight_layout()
        p_l = os.path.join(output_dir, '_chart_rebate_low.png')
        fig2.savefig(p_l, dpi=150, bbox_inches='tight')
        plt.close(fig2)
        paths['rebate_low'] = p_l

    return paths


def img_to_base64(path):
    with open(path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')


def build_note_interaction_trend_data(note_data):
    """从笔记分析数据提取内容表现分析 + 每日互动趋势"""
    content = note_data.get('result', {}).get('content', {})
    overall = content.get('overall', {})
    interaction = overall.get('interaction_analysis', {})
    time_trend = overall.get('time_trend', {})

    daily_trend = []
    for item in time_trend.get('daily_trend', []):
        if not isinstance(item, dict):
            continue
        daily_trend.append({
            'date': _safe_str(item.get('日期', '')),
            'total_interaction': float(item.get('总互动量', 0) or 0),
            'avg_interaction': float(item.get('平均互动量', 0) or 0),
            'note_count': int(float(item.get('笔记数', 0) or 0)),
            'total_read': float(item.get('总阅读量', 0) or 0),
        })

    stats = {}
    raw_stats = interaction.get('stats', {})
    if isinstance(raw_stats, dict):
        stats = {
            'total_interaction': _safe_str(raw_stats.get('总互动量', '-')),
            'avg_interaction': _safe_str(raw_stats.get('平均互动量', '-')),
            'median_interaction': _safe_str(raw_stats.get('中位数互动量', '-')),
            'max_interaction': _safe_str(raw_stats.get('最大互动量', '-')),
            'avg_interaction_rate': _safe_str(raw_stats.get('平均互动率(%)', '-')),
            'median_interaction_rate': _safe_str(raw_stats.get('中位数互动率(%)', '-')),
        }

    group_analysis = interaction.get('group_analysis', [])

    return {
        'daily_trend': daily_trend,
        'stats': stats,
        'group_analysis': group_analysis,
    }


def generate_daily_interaction_chart(trend_data, output_dir):
    """生成每日互动量趋势折线图（含标注）"""
    paths = {}
    if not HAS_MPL or not trend_data.get('daily_trend'):
        return paths

    daily = trend_data['daily_trend']
    if not daily:
        return paths

    font_name = _try_get_font()
    if font_name:
        plt.rcParams['font.family'] = font_name
    plt.rcParams['axes.unicode_minus'] = False

    dates = [d['date'] for d in daily]
    interactions = [d['total_interaction'] for d in daily]
    read_counts = [d['total_read'] for d in daily]

    fig, ax1 = plt.subplots(figsize=(10, 5))

    # 主坐标轴：总互动量折线
    x = list(range(len(dates)))
    ax1.plot(x, interactions, color='#4361ee', marker='o', linewidth=2, markersize=4, label='总互动量', zorder=3)
    ax1.fill_between(x, interactions, alpha=0.08, color='#4361ee')

    # 标注最大值
    max_idx = interactions.index(max(interactions))
    max_val = interactions[max_idx]
    ax1.annotate(f'峰值: {max_val:.0f}\n({dates[max_idx]})',
                 xy=(max_idx, max_val), xytext=(max_idx, max_val + max(interactions) * 0.15),
                 ha='center', fontsize=9, color='#e07c24', fontweight='bold',
                 fontfamily=font_name,
                 arrowprops=dict(arrowstyle='->', color='#e07c24', lw=1.2),
                 bbox=dict(boxstyle='round,pad=0.3', facecolor='#fff3e0', edgecolor='#e07c24', alpha=0.9))

    # 均值线
    avg_val = sum(interactions) / len(interactions)
    ax1.axhline(avg_val, color='#ff6b6b', linestyle='--', linewidth=1, alpha=0.6, label=f'均值 ({avg_val:.0f})')
    ax1.text(len(dates) - 1, avg_val + max(interactions) * 0.02, f'均值: {avg_val:.0f}',
             fontsize=8, color='#ff6b6b', ha='right', fontfamily=font_name)

    # 3日移动平均线
    if len(daily) >= 3:
        ma3 = []
        for i in range(len(daily)):
            start = max(0, i - 2)
            segment = interactions[start:i + 1]
            ma3.append(sum(segment) / len(segment))
        ax1.plot(x, ma3, color='#2d6a4f', linewidth=1.5, linestyle=':', alpha=0.8, label='3日移动平均', zorder=2)

    ax1.set_xticks(x)
    ax1.set_xticklabels(dates, rotation=45, ha='right', fontsize=8)
    ax1.set_ylabel('互动量', fontsize=11, fontfamily=font_name)
    ax1.set_title('每日互动量趋势', fontsize=14, fontweight='bold', fontfamily=font_name)
    ax1.legend(fontsize=9, loc='upper left')
    ax1.tick_params(axis='y', labelsize=9)
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)

    fig.tight_layout()
    p = os.path.join(output_dir, '_chart_daily_interaction.png')
    fig.savefig(p, dpi=150, bbox_inches='tight')
    plt.close(fig)
    paths['daily_interaction'] = p

    return paths


def build_note_viral_data(note_data):
    """从笔记分析提取爆款笔记数据，按达人量级分组"""
    viral = note_data.get('result', {}).get('content', {}).get('overall', {}).get('viral_notes', [])
    if not viral:
        return {}
    grouped = {}
    for item in viral:
        if not isinstance(item, dict):
            continue
        tier = str(item.get('达人量级', ''))
        if not tier:
            continue
        if tier not in grouped:
            grouped[tier] = []
        def safe_str(v):
            if v is None:
                return ''
            return str(v)
        def safe_num(v):
            if v is None:
                return 0
            try:
                return float(v)
            except (ValueError, TypeError):
                return 0
        grouped[tier].append({
            'name': safe_str(item.get('达人昵称', '')),
            'project': safe_str(item.get('项目名称', '')),
            'tier': tier,
            'note_type': safe_str(item.get('笔记类型', '')),
            'interaction': int(safe_num(item.get('互动量', 0))),
            'read_count': int(safe_num(item.get('阅读量', 0))),
            'exposure': int(safe_num(item.get('曝光量', 0))),
            'interaction_rate': safe_str(item.get('互动率(%)', '')),
            'cpm': safe_str(item.get('cpm', '')),
            'cpe': safe_str(item.get('cpe', '')),
            'publish_time': safe_str(item.get('发布时间', '')),
            'pgy_url': safe_str(item.get('pgy_url', '')),
            'home_url': safe_str(item.get('home_url', '')),
            'note_url': safe_str(item.get('note_url', '')),
        })
    # 每组按互动量降序
    for tier in grouped:
        grouped[tier].sort(key=lambda r: r['interaction'], reverse=True)
    return grouped


def _build_cost_lookup(note_data):
    """构建 达人昵称 -> cost_amount 查找表"""
    by_tier = note_data.get('result', {}).get('content', {}).get('by_tier', {})
    lookup = {}
    for tkey in ['KOC', 'KOL', '十万KOL']:
        tier_data = by_tier.get(tkey, {})
        if not isinstance(tier_data, dict):
            continue
        for item in tier_data.get('interaction_analysis', {}).get('top_interaction_notes', []):
            if not isinstance(item, dict):
                continue
            name = str(item.get('influencer_nickname', ''))
            cost = item.get('cost_amount', 0)
            if name and name not in lookup:
                lookup[name] = float(cost) if cost else 0
    return lookup


def build_note_high_value_data(note_data):
    """从笔记分析提取高价值达人数据（按 tier 分组，合并成本）"""
    overall_viral = note_data.get('result', {}).get('content', {}).get('overall', {}).get('viral_notes', [])
    cost_lookup = _build_cost_lookup(note_data)

    tier_map = {}
    for tkey in ['KOC', 'KOL', '十万KOL']:
        items = []
        for item in overall_viral:
            if not isinstance(item, dict):
                continue
            tier = str(item.get('达人量级', ''))
            if tier != tkey:
                continue
            name = str(item.get('达人昵称', ''))
            items.append({
                'name': name,
                'tier': tier,
                'interaction': int(float(item.get('互动量', 0) or 0)),
                'cost': cost_lookup.get(name, 0),
                'cpe': str(item.get('cpe', '')),
                'interaction_rate': str(item.get('互动率(%)', '')),
                'pgy_url': str(item.get('pgy_url', '')),
                'home_url': str(item.get('home_url', '')),
                'note_url': str(item.get('note_url', '')),
            })
        if items:
            items.sort(key=lambda r: r['interaction'], reverse=True)
            tier_map[tkey] = items
    return tier_map


def build_note_project_comparison_data(note_data):
    """从笔记分析提取项目效果对比（按 项目名称+达人量级 聚合）"""
    overall_viral = note_data.get('result', {}).get('content', {}).get('overall', {}).get('viral_notes', [])
    cost_lookup = _build_cost_lookup(note_data)

    groups = {}
    for item in overall_viral:
        if not isinstance(item, dict):
            continue
        project = str(item.get('项目名称', ''))
        tier = str(item.get('达人量级', ''))
        if not project or not tier:
            continue
        key = (project, tier)
        if key not in groups:
            groups[key] = {
                'project': project, 'tier': tier,
                'interactions': [], 'reads': [], 'costs': [],
                'cpms': [], 'cpes': [], 'rates': [],
            }
        g = groups[key]
        g['interactions'].append(float(item.get('互动量', 0) or 0))
        g['reads'].append(float(item.get('阅读量', 0) or 0))
        name = str(item.get('达人昵称', ''))
        g['costs'].append(cost_lookup.get(name, 0))
        g['cpms'].append(float(item.get('cpm', 0) or 0))
        g['cpes'].append(float(item.get('cpe', 0) or 0))
        try:
            r = float(str(item.get('互动率(%)', '0')).replace('%', ''))
        except (ValueError, TypeError):
            r = 0
        g['rates'].append(r)

    rows = []
    for key, g in groups.items():
        n = len(g['interactions'])
        avg_int = sum(g['interactions']) / n if n else 0
        avg_read = sum(g['reads']) / n if n else 0
        avg_cost = sum(g['costs']) / n if n else 0
        avg_cpm = sum(g['cpms']) / n if n else 0
        avg_cpe = sum(g['cpes']) / n if n else 0
        avg_rate = sum(g['rates']) / n if n else 0

        # 综合得分：基于互动量、互动率、CPE效率的综合评估
        score = 0
        score += min(avg_int / 500, 30)  # 互动量得分（最高30）
        score += min(avg_rate * 1.5, 25)  # 互动率得分（最高25）
        if avg_cpe > 0:
            cpe_score = min(100 / avg_cpe, 25)  # CPE越低分越高（最高25）
        else:
            cpe_score = 20
        score += cpe_score
        score += min(avg_read / 5000 * 20, 20)  # 阅读量得分（最高20）
        score = min(round(score, 2), 100)

        if score >= 85:
            rating = '优秀'
        elif score >= 70:
            rating = '良好'
        elif score >= 50:
            rating = '一般'
        else:
            rating = '较差'

        rows.append({
            'project': g['project'],
            'tier': g['tier'],
            'avg_interaction': round(avg_int, 0),
            'avg_read': round(avg_read, 0),
            'avg_cost': round(avg_cost, 2),
            'avg_cpm': round(avg_cpm, 2),
            'avg_cpe': round(avg_cpe, 2),
            'avg_rate': round(avg_rate, 2),
            'score': score,
            'rating': rating,
        })

    rows.sort(key=lambda r: (r['rating'] == '优秀' and 0 or r['rating'] == '良好' and 1 or r['rating'] == '一般' and 2 or 3,
                             -r['score']))
    return rows


def generate_weekly_report(analysis_id, source='analysis_results', week_label=None, report_type='weekly', note_analysis_id=''):
    data = load_analysis_data(analysis_id, source) if analysis_id else {}
    is_monthly = report_type == 'monthly'
    if week_label is None:
        week_label = calc_month_label() if is_monthly else calc_week_label()
    type_name = '月报' if is_monthly else '周报'

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'媒介-审计报告（{week_label}）')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.bold = True

    doc.add_paragraph('')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f'分析ID: {analysis_id}    生成日期: {datetime.now().strftime("%Y-%m-%d")}')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('').paragraph_format.space_after = Pt(12)

    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs', 'reports')
    os.makedirs(output_dir, exist_ok=True)

    h1 = doc.add_heading('一、总体分析', level=1)
    for r in h1.runs:
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    h2 = doc.add_heading('1、总结', level=2)
    for r in h2.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

    # ── 总结表格 ──
    table_rows = build_summary_table_data(data)

    headers = ['所属小组', '媒介人数', '签框量(口头)', '签框量(书面)', '总提报量', '总定档量', '平均定档数', '过筛率中位数', '平均定档成本']
    table = doc.add_table(rows=1 + len(table_rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _add_cell_text(cell, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_shading(cell, '4361ee')

    for row_idx, row_data in enumerate(table_rows):
        vals = [
            row_data['group'],
            str(row_data['media_count']),
            str(row_data['kou']),
            str(row_data['shu']),
            str(row_data['tibao']),
            str(row_data['dingdang']),
            str(row_data['avg_dingdang']),
            row_data['guoshuai'],
            f"{row_data['avg_cost']:.2f}" if isinstance(row_data['avg_cost'], (int, float)) else str(row_data['avg_cost'])
        ]
        for col_idx, val in enumerate(vals):
            _add_cell_text(table.rows[row_idx + 1].cells[col_idx], val, size=9)
        # group color in first col
        gname = row_data['group']
        if gname in GROUP_COLORS:
            _set_cell_shading(table.rows[row_idx + 1].cells[0], GROUP_COLORS[gname][1])
        _set_alt_row_shading(table, row_idx + 1, start_col=1)

    # ── 2、工作量 ──
    doc.add_paragraph('')
    h_workload = doc.add_heading('2、工作量', level=2)
    for r in h_workload.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

    note = doc.add_paragraph()
    note_run = note.add_run('注：离职人员已去除')
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    note_run.italic = True

    detail_groups = build_workload_detail_data(data)

    if detail_groups:
        group_count = len(detail_groups)
        cols_per_group = 4  # 小组, 姓名, 提报数, 定档数
        total_cols = group_count * cols_per_group
        max_rows = max(len(g['people']) for g in detail_groups)

        table2 = doc.add_table(rows=2 + max_rows, cols=total_cols)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.style = 'Table Grid'

        # 第一行: 小组名称（合并 4 列）
        for gi, gd in enumerate(detail_groups):
            start_col = gi * cols_per_group
            gkey = gd['group']
            gcolor = GROUP_COLORS.get(gkey, ('#4361ee', '#dce6ff', '#f0f4ff'))[0]
            cell = table2.rows[0].cells[start_col]
            _add_cell_text(cell, gkey, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
            _set_cell_shading(cell, gcolor)
            for ci in range(start_col + 1, start_col + cols_per_group):
                cell = table2.rows[0].cells[ci]
                _set_cell_shading(cell, gcolor)
                table2.rows[0].cells[start_col].merge(cell)

        # 第二行: 列标题
        sub_headers = ['姓名', '提报数', '定档数']
        for gi in range(group_count):
            start_col = gi * cols_per_group
            gkey = detail_groups[gi]['group']
            gbg_lighter = GROUP_COLORS.get(gkey, ('#4361ee', '#e8edff', '#f0f4ff'))[2]
            _add_cell_text(table2.rows[1].cells[start_col], '小组', bold=True, size=8)
            _set_cell_shading(table2.rows[1].cells[start_col], gbg_lighter)
            for si, sh in enumerate(sub_headers):
                cell = table2.rows[1].cells[start_col + 1 + si]
                _add_cell_text(cell, sh, bold=True, size=8)
                _set_cell_shading(cell, gbg_lighter)

        # 数据行
        for ri in range(max_rows):
            for gi, gd in enumerate(detail_groups):
                start_col = gi * cols_per_group
                gkey = gd['group']
                gbg_mid = GROUP_COLORS.get(gkey, ('#4361ee', '#f0f4ff', '#f0f4ff'))[1]
                gbg_alt = GROUP_COLORS.get(gkey, ('#4361ee', '#f8f9fa', '#f8f9fa'))[2]
                if ri < len(gd['people']):
                    p = gd['people'][ri]
                    bg = gbg_mid if ri % 2 == 0 else gbg_alt
                    for ci in range(cols_per_group):
                        _set_cell_shading(table2.rows[2 + ri].cells[start_col + ci], bg)
                    gc = GROUP_COLORS.get(gkey, ('#4361ee', '', ''))[0]
                    _add_cell_text(table2.rows[2 + ri].cells[start_col], gkey, size=8, bold=True,
                                   color=_hex_to_rgb(gc))
                    _add_cell_text(table2.rows[2 + ri].cells[start_col + 1], p['name'], size=8)
                    _add_cell_text(table2.rows[2 + ri].cells[start_col + 2], str(p['tibao']), size=8)
                    _add_cell_text(table2.rows[2 + ri].cells[start_col + 3], str(p['dingdang']), size=8)
                else:
                    for ci in range(cols_per_group):
                        _add_cell_text(table2.rows[2 + ri].cells[start_col + ci], '', size=8)

    chart_data = build_chart_data(data)
    if chart_data and HAS_MPL:
        report_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'outputs', 'reports')
        os.makedirs(report_dir, exist_ok=True)
        chart_paths = generate_chart_images(chart_data, report_dir)

        if 'guoshuai' in chart_paths:
            p_gs = doc.add_paragraph()
            p_gs.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_gs = p_gs.add_run('图1: 各小组过筛率中位数')
            run_gs.font.size = Pt(10)
            run_gs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_picture(chart_paths['guoshuai'], width=Cm(14))

        if 'dingdang' in chart_paths:
            doc.add_paragraph('')
            p_dd = doc.add_paragraph()
            p_dd.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_dd = p_dd.add_run('图2: 各小组总定档量')
            run_dd.font.size = Pt(10)
            run_dd.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_picture(chart_paths['dingdang'], width=Cm(14))

    # ── 3、工作质量 ──
    doc.add_paragraph('')
    h_qual = doc.add_heading('3、工作质量', level=2)
    for r in h_qual.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

    qual_note = doc.add_paragraph()
    qn_run = qual_note.add_run('注：由于各组的表头字段不统一，且与数据库字段不匹配，因此过筛滤以最终通过的过筛为准。')
    qn_run.font.size = Pt(9)
    qn_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    qn_run.italic = True

    premium_people = build_quality_detail_data(data, 'premium')
    high_read_people = build_quality_detail_data(data, 'high_read')

    if premium_people:
        sub_h1 = doc.add_heading('（1）优质达人', level=3)
        for r in sub_h1.runs:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        q_headers = ['对应名字', '所属小组', '总提报达人数', '过筛人数', '过筛率(%)', '质量评估']
        q_table = doc.add_table(rows=1 + len(premium_people), cols=len(q_headers))
        q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        q_table.style = 'Table Grid'
        for i, h in enumerate(q_headers):
            _add_cell_text(q_table.rows[0].cells[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
            _set_cell_shading(q_table.rows[0].cells[i], 'EF5350')
        for ri, p in enumerate(premium_people):
            vals = [p['name'], p['group'], str(p['total']), str(p['passed']), p['rate_str'], p['evaluation']]
            for ci, v in enumerate(vals):
                _add_cell_text(q_table.rows[ri + 1].cells[ci], v, size=8)
            _set_cell_shading(q_table.rows[ri + 1].cells[0], 'FFEBEE')
            _set_alt_row_shading(q_table, ri + 1)

    if high_read_people:
        doc.add_paragraph('')
        sub_h2 = doc.add_heading('（2）高阅读达人', level=3)
        for r in sub_h2.runs:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        q2_headers = ['对应名字', '所属小组', '总提报达人数', '过筛人数', '过筛率', '过筛率(%)', '质量评估']
        q2_table = doc.add_table(rows=1 + len(high_read_people), cols=len(q2_headers))
        q2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        q2_table.style = 'Table Grid'
        for i, h in enumerate(q2_headers):
            _add_cell_text(q2_table.rows[0].cells[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
            _set_cell_shading(q2_table.rows[0].cells[i], 'EF5350')
        for ri, p in enumerate(high_read_people):
            vals = [p['name'], p['group'], str(p['total']), str(p['passed']),
                    '%.2f' % p['rate_decimal'], p['rate_str'], p['evaluation']]
            for ci, v in enumerate(vals):
                _add_cell_text(q2_table.rows[ri + 1].cells[ci], v, size=8)
            _set_cell_shading(q2_table.rows[ri + 1].cells[0], 'FFEBEE')
            _set_alt_row_shading(q2_table, ri + 1)

    if premium_people or high_read_people:
        doc.add_paragraph('')
        legend = doc.add_paragraph()
        lr = legend.add_run('注：\n- 优秀：80%及以上\n- 良好：65%-80%\n- 一般：50%-64%\n- 待改进：40%-49%\n- 较差：低于40%')
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 4、成本发挥 ──
    doc.add_paragraph('')
    h_cost = doc.add_heading('4、成本发挥', level=2)
    for r in h_cost.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

    cost_perf = build_cost_performance_data(data)
    if cost_perf:
        c_headers = ['媒介组', '定档达人数', '总项目数', '平均成本(元)', '平均下单价(元)', '平均返点比例(%)', '返点比例最高', '返点比例最低']
        c_table = doc.add_table(rows=1 + len(cost_perf), cols=len(c_headers))
        c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_table.style = 'Table Grid'
        for i, h in enumerate(c_headers):
            _add_cell_text(c_table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
            _set_cell_shading(c_table.rows[0].cells[i], '42A5F5')
        for ri, r in enumerate(cost_perf):
            vals = [r['group'], str(r['dingdang']), str(r['total_project']), '%.2f' % r['avg_cost'] if isinstance(r['avg_cost'], (int, float)) else str(r['avg_cost']), '%.2f' % r['avg_order'] if isinstance(r['avg_order'], (int, float)) else str(r['avg_order']), r['avg_rebate'], r['high'], r['low']]
            for ci, v in enumerate(vals):
                _add_cell_text(c_table.rows[ri + 1].cells[ci], v, size=9)
            _set_cell_shading(c_table.rows[ri + 1].cells[0], 'E3F2FD')

        doc.add_paragraph('')
        c_note = doc.add_paragraph()
        cn_run = c_note.add_run('注：不含高返项目')
        cn_run.font.size = Pt(8)
        cn_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        cn_run.italic = True

        # ── 返点比例分布图 ──
        if HAS_MPL:
            reb_paths = generate_rebate_chart_images(data, output_dir)
            if 'rebate_high' in reb_paths:
                doc.add_paragraph('')
                p_rh = doc.add_paragraph()
                p_rh.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_rh = p_rh.add_run('图3: 高返点媒介分布（平均返点比例 >35%）')
                run_rh.font.size = Pt(10)
                run_rh.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_picture(reb_paths['rebate_high'], width=Cm(14))
            if 'rebate_low' in reb_paths:
                doc.add_paragraph('')
                p_rl = doc.add_paragraph()
                p_rl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_rl = p_rl.add_run('图4: 低返点媒介分布（平均返点比例 <20%）')
                run_rl.font.size = Pt(10)
                run_rl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_picture(reb_paths['rebate_low'], width=Cm(14))

        # ── （1）返点分析 ──
        doc.add_paragraph('')
        sub_r1 = doc.add_heading('（1）返点分析', level=3)
        for r in sub_r1.runs:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        sub_r1a = doc.add_heading('①返点整体效益分析', level=4)
        for r in sub_r1a.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

        reb_rows = build_rebate_analysis_data(data)
        if reb_rows:
            r_headers = ['定档媒介', '定档达人数', '所属小组', '平均返点比例(%)', '返点比例最大值(%)',
                         '返点比例最小值(%)', '返点比例中位数(%)', '总返点金额(元)', '平均返点金额(元)',
                         '返点金额最大值(元)', '返点金额最小值(元)', '返点金额中位数(元)',
                         '总下单价(元)', '平均下单价(元)', '返点表现评估', '返点优化建议']
            r_table = doc.add_table(rows=1 + len(reb_rows), cols=len(r_headers))
            r_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            r_table.style = 'Table Grid'
            r_table.autofit = True
            font_sz = 7
            for i, h in enumerate(r_headers):
                _add_cell_text(r_table.rows[0].cells[i], h, bold=True, size=font_sz,
                               color=RGBColor(0xFF, 0xFF, 0xFF))
                _set_cell_shading(r_table.rows[0].cells[i], '42A5F5')
            for ri, rr in enumerate(reb_rows):
                vals = [
                    rr['name'],
                    str(rr['dingdang']),
                    rr['group'],
                    rr['avg_rebate'],
                    rr['max_rebate'],
                    rr['min_rebate'],
                    rr['median_rebate'],
                    '%.2f' % rr['total_rebate_amt'],
                    '%.2f' % rr['avg_rebate_amt'],
                    '%.2f' % rr['max_rebate_amt'],
                    '%.2f' % rr['min_rebate_amt'],
                    '%.2f' % rr['median_rebate_amt'],
                    '%.2f' % rr['total_order'] if isinstance(rr['total_order'], (int, float)) else str(rr['total_order']),
                    '%.2f' % rr['avg_order'] if isinstance(rr['avg_order'], (int, float)) else str(rr['avg_order']),
                    rr['evaluation'],
                    rr['suggestion']
                ]
                for ci, v in enumerate(vals):
                    _add_cell_text(r_table.rows[ri + 1].cells[ci], v, size=font_sz)
            for ri, rr in enumerate(reb_rows):
                _set_cell_shading(r_table.rows[ri + 1].cells[0], 'E3F2FD')
                _set_alt_row_shading(r_table, ri + 1)
                ev_cell = r_table.rows[ri + 1].cells[14]
                ev = rr['evaluation']
                if ev == '优秀':
                    _set_cell_shading(ev_cell, 'd4edda')
                elif ev == '较差':
                    _set_cell_shading(ev_cell, 'f8d7da')
                elif ev == '一般':
                    _set_cell_shading(ev_cell, 'fff3cd')

        # ── ②基于达人量级分析 ──
        sub_r1b = doc.add_heading('②基于达人量级分析', level=4)
        for r in sub_r1b.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

        lvl_rows = build_level_analysis_data(data)
        if lvl_rows:
            l_headers = ['定档媒介', '达人量级', '达人数', '所属小组', '总成本(元)', '平均成本(元)',
                         '总下单价(元)', '平均下单价(元)', '总返点金额(元)', '平均返点金额(元)',
                         '平均返点比例(%)', '总互动量', '平均互动量', '总阅读量', '平均阅读量',
                         '平均CPE', '平均CPM']
            l_table = doc.add_table(rows=1 + len(lvl_rows), cols=len(l_headers))
            l_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            l_table.style = 'Table Grid'
            l_table.autofit = True
            fs = 7
            for i, h in enumerate(l_headers):
                _add_cell_text(l_table.rows[0].cells[i], h, bold=True, size=fs,
                               color=RGBColor(0xFF, 0xFF, 0xFF))
                _set_cell_shading(l_table.rows[0].cells[i], '42A5F5')
            for ri, rr in enumerate(lvl_rows):
                vals = [
                    rr['name'], rr['level'], str(rr['count']), rr['group'],
                    '%.2f' % rr['total_cost'], '%.2f' % rr['avg_cost'],
                    '%.2f' % rr['total_order'] if isinstance(rr['total_order'], (int, float)) else str(rr['total_order']),
                    '%.2f' % rr['avg_order'] if isinstance(rr['avg_order'], (int, float)) else str(rr['avg_order']),
                    '%.2f' % rr['total_rebate'], '%.2f' % rr['avg_rebate'],
                    rr['avg_rebate_pct'],
                    '%.0f' % rr['total_interact'], '%.2f' % rr['avg_interact'],
                    '%.0f' % rr['total_read'], '%.2f' % rr['avg_read'],
                    '%.2f' % rr['avg_cpe'], '%.2f' % rr['avg_cpm']
                ]
                for ci, v in enumerate(vals):
                    _add_cell_text(l_table.rows[ri + 1].cells[ci], v, size=fs)
                _set_cell_shading(l_table.rows[ri + 1].cells[0], 'E3F2FD')
                _set_alt_row_shading(l_table, ri + 1)

    # ── （2）成本分析 ──
    doc.add_paragraph('')
    sub_ca = doc.add_heading('（2）成本分析', level=3)
    for r in sub_ca.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub_ca1 = doc.add_heading('定档媒介成本分析', level=4)
    for r in sub_ca1.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

    cost_rows = build_cost_analysis_data(data)
    if cost_rows:
        ca_headers = ['定档媒介', '定档达人数', '所属小组', '总成本(元)', '平均成本(元)',
                      '成本最大值(元)', '成本最小值(元)', '成本中位数(元)', '总报价(元)',
                      '平均报价(元)', '报价最大值(元)', '报价最小值(元)', '总下单价(元)',
                      '平均下单价(元)', '总节约金额(元)', '平均节约金额(元)', '成本占比(%)',
                      '总返点金额(元)', '平均返点金额(元)']
        ca_table = doc.add_table(rows=1 + len(cost_rows), cols=len(ca_headers))
        ca_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ca_table.style = 'Table Grid'
        ca_table.autofit = True
        fs = 7
        for i, h in enumerate(ca_headers):
            _add_cell_text(ca_table.rows[0].cells[i], h, bold=True, size=fs,
                           color=RGBColor(0xFF, 0xFF, 0xFF))
            _set_cell_shading(ca_table.rows[0].cells[i], '42A5F5')
        for ri, rr in enumerate(cost_rows):
            vals = [
                rr['name'], str(rr['count']), rr['group'],
                '%.2f' % rr['total_cost'], '%.2f' % rr['avg_cost'],
                '%.2f' % rr['max_cost'], '%.2f' % rr['min_cost'],
                '%.2f' % rr['median_cost'],
                '%.2f' % rr['total_quote'], '%.2f' % rr['avg_quote'],
                '%.2f' % rr['max_quote'], '%.2f' % rr['min_quote'],
                '%.2f' % rr['total_order'], '%.2f' % rr['avg_order'],
                '%.2f' % rr['total_save'], '%.2f' % rr['avg_save'],
                rr['cost_pct'],
                '%.2f' % rr['total_rebate'], '%.2f' % rr['avg_rebate']
            ]
            for ci, v in enumerate(vals):
                _add_cell_text(ca_table.rows[ri + 1].cells[ci], v, size=fs)
            _set_cell_shading(ca_table.rows[ri + 1].cells[0], 'E3F2FD')
            _set_alt_row_shading(ca_table, ri + 1)

    # ── 二、笔记分析 ──
    if note_analysis_id:
        try:
            note_data = load_analysis_data(note_analysis_id, source='note_analysis_results')
        except Exception:
            note_data = None
        if note_data:
            analysis_types = note_data.get('analysis_types', [])
            has_content = 'content' in analysis_types
            has_value = 'value' in analysis_types
            has_review = 'review' in analysis_types

            if has_content or has_value or has_review:
                doc.add_paragraph('')
                h_note = doc.add_heading('二、笔记分析', level=1)
                for r in h_note.runs:
                    r.font.size = Pt(16)
                    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

            if has_content:
                trend_data = build_note_interaction_trend_data(note_data)
                stats = trend_data.get('stats', {})

                # 内容表现分析表
                if stats:
                    sub_note1 = doc.add_heading('内容表现分析', level=2)
                    for r in sub_note1.runs:
                        r.font.size = Pt(12)
                        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

                    stat_headers = ['总互动量', '平均互动量', '中位数互动量', '最大互动量', '平均互动率(%)', '中位数互动率(%)']
                    stat_keys = ['total_interaction', 'avg_interaction', 'median_interaction', 'max_interaction',
                                 'avg_interaction_rate', 'median_interaction_rate']
                    st_rows = [[stats.get(k, '-') for k in stat_keys]]
                    st_table = doc.add_table(rows=2, cols=len(stat_headers))
                    st_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    st_table.style = 'Table Grid'
                    st_table.autofit = True
                    fs = 10
                    for i, h in enumerate(stat_headers):
                        _add_cell_text(st_table.rows[0].cells[i], h, bold=True, size=fs,
                                       color=RGBColor(0xFF, 0xFF, 0xFF))
                        _set_cell_shading(st_table.rows[0].cells[i], '66BB6A')
                    for ci, v in enumerate(st_rows[0]):
                        _add_cell_text(st_table.rows[1].cells[ci], v, size=fs)
                    _set_cell_shading(st_table.rows[1].cells[0], 'E8F5E9')

                # 每日互动量趋势图
                if HAS_MPL:
                    chart_paths = generate_daily_interaction_chart(trend_data, output_dir)
                    if 'daily_interaction' in chart_paths:
                        doc.add_paragraph('')
                        sub_note2 = doc.add_heading('每日互动量趋势', level=2)
                        for r in sub_note2.runs:
                            r.font.size = Pt(12)
                            r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)
                        chart_p = doc.add_paragraph()
                        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        chart_p.add_run().add_picture(chart_paths['daily_interaction'], width=Inches(6.0))

                # 爆款笔记分析
                viral_data = build_note_viral_data(note_data)
                if viral_data:
                    doc.add_paragraph('')
                    sub_note3 = doc.add_heading('爆款笔记分析', level=2)
                    for r in sub_note3.runs:
                        r.font.size = Pt(12)
                        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

                    tier_order = ['KOC', 'KOL', '十万KOL']
                    viral_headers = ['达人昵称', '项目名称', '达人量级', '笔记类型', '互动量', '阅读量',
                                     '曝光量', '互动率(%)', 'cpm', 'cpe', '发布时间', '蒲公英链接',
                                     '主页链接', '笔记链接']
                    vfs = 6
                    for ti, tier in enumerate(tier_order):
                        if tier not in viral_data:
                            continue
                        tier_label = {'KOC': 'KOC', 'KOL': 'KOL', '十万KOL': '十万KOL'}.get(tier, tier)
                        sub_tier = doc.add_heading('1.%d 爆款笔记 TOP（%s）' % (ti + 1, tier_label), level=3)
                        for r in sub_tier.runs:
                            r.font.size = Pt(10)
                            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

                        notes = viral_data[tier]
                        v_table = doc.add_table(rows=1 + len(notes), cols=len(viral_headers))
                        v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        v_table.style = 'Table Grid'
                        v_table.autofit = True
                        for i, h in enumerate(viral_headers):
                            _add_cell_text(v_table.rows[0].cells[i], h, bold=True, size=vfs,
                                           color=RGBColor(0xFF, 0xFF, 0xFF))
                            _set_cell_shading(v_table.rows[0].cells[i], '66BB6A')
                        for ri, note in enumerate(notes):
                            vals = [
                                note['name'], note['project'], note['tier'], note['note_type'],
                                str(note['interaction']), str(note['read_count']),
                                str(note['exposure']), note['interaction_rate'],
                                note['cpm'], note['cpe'], note['publish_time'][:10] if note['publish_time'] else '',
                                note['pgy_url'], note['home_url'], note['note_url']
                            ]
                            for ci, v in enumerate(vals):
                                _add_cell_text(v_table.rows[ri + 1].cells[ci], v, size=vfs)
                            _set_cell_shading(v_table.rows[ri + 1].cells[0], 'E8F5E9')
                            _set_alt_row_shading(v_table, ri + 1)

            if has_value:
                hv_data = build_note_high_value_data(note_data)
                if hv_data:
                    doc.add_paragraph('')
                    sub_note4 = doc.add_heading('2.高价值达人分析', level=2)
                    for r in sub_note4.runs:
                        r.font.size = Pt(12)
                        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

                    hv_headers = ['达人昵称', '达人量级', '互动量', '成本', 'CPE', '互动率(%)',
                                  '蒲公英链接', '主页链接', '笔记链接']
                    hv_tier_order = ['KOC', 'KOL', '十万KOL']
                    hv_fs = 7
                    for hti, tier in enumerate(hv_tier_order):
                        if tier not in hv_data:
                            continue
                        tier_label = {'KOC': 'KOC', 'KOL': 'KOL', '十万KOL': '十万KOL'}.get(tier, tier)
                        sub_hv = doc.add_heading('2.%d 高价值达人（%s）' % (hti + 1, tier_label), level=3)
                        for r in sub_hv.runs:
                            r.font.size = Pt(10)
                            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

                        notes = hv_data[tier]
                        hv_table = doc.add_table(rows=1 + len(notes), cols=len(hv_headers))
                        hv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        hv_table.style = 'Table Grid'
                        hv_table.autofit = True
                        for i, h in enumerate(hv_headers):
                            _add_cell_text(hv_table.rows[0].cells[i], h, bold=True, size=hv_fs,
                                           color=RGBColor(0xFF, 0xFF, 0xFF))
                            _set_cell_shading(hv_table.rows[0].cells[i], '66BB6A')
                        for ri, note in enumerate(notes):
                            vals = [
                                note['name'], note['tier'], str(note['interaction']),
                                '%.2f' % note['cost'] if note['cost'] else '0',
                                note['cpe'], note['interaction_rate'],
                                note['pgy_url'], note['home_url'], note['note_url']
                            ]
                            for ci, v in enumerate(vals):
                                _add_cell_text(hv_table.rows[ri + 1].cells[ci], v, size=hv_fs)
                            _set_cell_shading(hv_table.rows[ri + 1].cells[0], 'E8F5E9')
                            _set_alt_row_shading(hv_table, ri + 1)

            if has_review:
                pc_rows = build_note_project_comparison_data(note_data)
                if pc_rows:
                    doc.add_paragraph('')
                    sub_note5 = doc.add_heading('3.项目效果对比', level=2)
                    for r in sub_note5.runs:
                        r.font.size = Pt(12)
                        r.font.color.rgb = RGBColor(0x43, 0x61, 0xee)

                    pc_headers = ['项目名称', '达人量级', '平均互动量', '平均阅读量', '平均成本',
                                  '平均CPM', '平均CPE', '平均互动率(%)', '综合得分', '效果评级']
                    pc_fs = 8
                    pc_table = doc.add_table(rows=1 + len(pc_rows), cols=len(pc_headers))
                    pc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    pc_table.style = 'Table Grid'
                    pc_table.autofit = True
                    for i, h in enumerate(pc_headers):
                        _add_cell_text(pc_table.rows[0].cells[i], h, bold=True, size=pc_fs,
                                       color=RGBColor(0xFF, 0xFF, 0xFF))
                        _set_cell_shading(pc_table.rows[0].cells[i], '66BB6A')
                    for ri, rr in enumerate(pc_rows):
                        vals = [
                            rr['project'], rr['tier'], '%.0f' % rr['avg_interaction'],
                            '%.0f' % rr['avg_read'], '%.2f' % rr['avg_cost'],
                            '%.2f' % rr['avg_cpm'], '%.2f' % rr['avg_cpe'],
                            '%.2f' % rr['avg_rate'], '%.2f' % rr['score'], rr['rating']
                        ]
                        for ci, v in enumerate(vals):
                            _add_cell_text(pc_table.rows[ri + 1].cells[ci], v, size=pc_fs)
                        _set_cell_shading(pc_table.rows[ri + 1].cells[0], 'E8F5E9')
                        _set_alt_row_shading(pc_table, ri + 1)
                        ev = rr['rating']
                        ev_cell = pc_table.rows[ri + 1].cells[-1]
                        if ev == '优秀':
                            _set_cell_shading(ev_cell, 'd4edda')
                        elif ev == '良好':
                            _set_cell_shading(ev_cell, 'cfe2ff')
                        elif ev == '一般':
                            _set_cell_shading(ev_cell, 'fff3cd')

    filename = f"媒介-审计报告（{week_label}）.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    logger.info(f"{type_name}生成成功: {filepath}")

    return filepath, filename


def generate_monthly_report(analysis_id='', source='analysis_results', month_label=None, note_analysis_id=''):
    return generate_weekly_report(analysis_id, source, week_label=month_label, report_type='monthly',
                                  note_analysis_id=note_analysis_id)
